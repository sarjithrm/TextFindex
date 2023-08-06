import argparse, re, ctypes, os, docx, PyPDF2
from docx import Document
import pandas as pd

extensions = ['txt','csv','xlsx','pdf','docx','xls']
returnType = 'line'

parser = argparse.ArgumentParser(
    prog='TextFinder',
    description='Scan the files of provided types/extension in provided path or system and return the ouput'
)

parser.add_argument('-f', '--fileExtensions', choices=extensions, nargs='+', help='List of file extensions to scan')
parser.add_argument('-r', '--returnType', choices=['line','paragraph'], default=returnType, help='Return type object sentence or paragraph')
parser.add_argument('-o', '--output', required=True, help='Output file path')
parser.add_argument('-i', '--input', nargs='+', help='Please provide the file or directory path to search')
parser.add_argument('-s', '--search', required=True, help='Search text')

args = parser.parse_args()

def is_ignored_path(path):
    ignored_dirs = [
        "Program Files",
        "Program Files (x86)",
        "Windows",
        "AppData",
    ]

    for ignored_dir in ignored_dirs:
        if ignored_dir.lower() in path.lower():
            return True
    return False

def get_all_drives():
    drives = []
    bitmask = ctypes.windll.kernel32.GetLogicalDrives()

    for letter in range(65, 91):
        mask = 1 << (letter - 65)
        if bitmask & mask:
            drives.append(chr(letter) + ':')

    return drives
    
def find_text_with_string(file_path, target_string, search_type):
    found_text = set()
    file_extension = file_path.split('.')[-1].lower()

    try:
        if file_extension == 'txt':
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                if search_type == 'line':
                    for line in text.split('\n'):
                        if re.search(r'\b' + re.escape(target_string) + r'\b', line, re.IGNORECASE):
                            found_text.add(line.strip())
                elif search_type == 'paragraph':
                    paragraphs = re.split(r'\n\s*\n', text)
                    for paragraph in paragraphs:
                        if re.search(r'\b' + re.escape(target_string) + r'\b', paragraph, re.IGNORECASE):
                            found_text.add(paragraph.strip())

        elif file_extension == 'docx':
            doc = docx.Document(file_path)
            if search_type == 'line':
                for paragraph in doc.paragraphs:
                    for line in paragraph.text.split('\n'):
                        if re.search(r'\b' + re.escape(target_string) + r'\b', line, re.IGNORECASE):
                            found_text.add(line.strip())
            elif search_type == 'paragraph':
                for paragraph in doc.paragraphs:
                    if re.search(r'\b' + re.escape(target_string) + r'\b', paragraph, re.IGNORECASE):
                        found_text.add(paragraph.text.strip())

        elif file_extension == 'pdf':
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfFileReader(file)
                for page_num in range(pdf_reader.getNumPages()):
                    page = pdf_reader.getPage(page_num)
                    page_content = page.extractText()
                    if search_type == 'line':
                        for line in page_content.split('\n'):
                            if re.search(r'\b' + re.escape(target_string) + r'\b', line, re.IGNORECASE):
                                found_text.add(line.strip())
                    elif search_type == 'paragraph':
                        paragraphs = re.split(r'\n\s*\n', page_content)
                        for paragraph in paragraphs:
                            if re.search(r'\b' + re.escape(target_string) + r'\b', paragraph, re.IGNORECASE):
                                found_text.add(paragraph.strip())

    except Exception as e:
        x = e

    return found_text

def find_rows_columns_with_string(file_path, target_string):
    found_rows_columns = set()
    file_extension = file_path.split('.')[-1].lower()

    try:
        if file_extension in ['csv', 'xls', 'xlsx']:
            df = pd.read_csv(file_path) if file_extension == 'csv' else pd.read_excel(file_path)

            for index, row in df.iterrows():
                for column, value in row.iteritems():
                    if pd.notna(value) and re.search(r'\b' + re.escape(target_string) + r'\b', str(value), re.IGNORECASE): 
                        found_rows_columns.add((index, column, str(value)))

    except Exception as e:
        x = e

    return found_rows_columns

def is_file_path(path):
    return os.path.isfile(path)

def process_single_file(file_path, target_string, search_type, extensions):
    results = set()
    file_extension = file_path.split('.')[-1].lower()
    if file_extension in extensions:
        if file_extension in ['txt', 'docx', 'pdf']:
            try:
                text = find_text_with_string(file_path, target_string, search_type)
                if text:
                    for t in text:
                        results.add((os.path.basename(file_path), 'N/A', 'N/A', t))
            except Exception as e:
                print(f"Error occurred while processing {file_path}: {e}")
        elif file_extension in ['csv', 'xls', 'xlsx']:
            try:
                rows_columns = find_rows_columns_with_string(file_path, target_string)
                if rows_columns:
                    for row_column in rows_columns:
                        results.add((os.path.basename(file_path), row_column[0], row_column[1], row_column[2]))
            except Exception as e:
                x = e
    return results

def scan_files_for_target(directory_path, target_string, extensions, search_type):
    results = set()

    if is_file_path(directory_path):
        results.update(process_single_file(directory_path, target_string, search_type, extensions))
    else:
        for root, _, files in os.walk(directory_path):
            if is_ignored_path(root):
                continue
            
            for file in files:
                file_path = os.path.join(root, file)
                results.update(process_single_file(file_path, target_string, search_type, extensions))

    return results

def write_to_excel(output_file, results):
    if not os.path.exists(output_file):
        with pd.ExcelWriter(output_file) as writer:
            df = pd.DataFrame(results, columns=['File Name', 'Row', 'Column', 'Content'])
            df.to_excel(writer, index=False)
    else:
        df = pd.read_excel(output_file)
        df = pd.concat([df, pd.DataFrame(results, columns=['File Name', 'Row', 'Column', 'Content'])], ignore_index=True)
        df.to_excel(output_file, index=False)
 
if args.fileExtensions is not None:
    extensions = args.fileExtensions
returnType = args.returnType
outputfile = args.output
searchPaths = get_all_drives() if args.input is None else args.input
search = args.search

for path in searchPaths:
    results = scan_files_for_target(path, search, extensions, returnType)

write_to_excel(outputfile, results)